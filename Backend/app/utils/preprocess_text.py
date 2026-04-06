# import io, re
# from PyPDF2 import PdfReader
# from docx import Document

# def extract_text(file_bytes, filename):
#     text = ""
#     if filename.endswith(".pdf"):
#         reader = PdfReader(io.BytesIO(file_bytes))
#         for page_num, page in enumerate(reader.pages):
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#             else:
#                 print(f"⚠️ Page {page_num} has no extractable text.")
    
#     elif filename.endswith(".docx"):
#         doc = Document(io.BytesIO(file_bytes))
#         for para in doc.paragraphs:
#             text += para.text + "\n"
    
#     elif filename.endswith(".txt"):
#         text = file_bytes.decode("utf-8")
    
#     else:
#         raise ValueError("Unsupported file type")
    
#     return text

# def clean_text(text):
#     text = re.sub(r'\s+', ' ', text)  # normalize spaces/newlines
#     return text.strip()

# def chunk_text(text, chunk_size=300):
#     words = text.split()
#     print(f"Total words: {len(words)}")
#     chunks = []
#     for i in range(0, len(words), chunk_size):
#         chunk = " ".join(words[i:i+chunk_size])
#         chunks.append(chunk)
#     print(f"✅ Total chunks created: {len(chunks)}")
#     return chunks

import io, os, re
from PyPDF2 import PdfReader
from docx import Document
from pdf2image import convert_from_bytes
import pytesseract

def extract_text(file_bytes, filename):
    text = ""
    if filename.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(file_bytes))
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
            else:
                print(f"⚠️ Page {page_num} has no extractable text. Using OCR...")
                # Convert this page to image and extract text with OCR
                images = convert_from_bytes(file_bytes, first_page=page_num+1, last_page=page_num+1)
                for image in images:
                    text += pytesseract.image_to_string(image, lang="eng") + "\n"

    elif filename.endswith(".docx"):
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"

    elif filename.endswith(".txt"):
        text = file_bytes.decode("utf-8")

    else:
        raise ValueError("Unsupported file type")

    return text


def clean_text(text):
    text = re.sub(r'\s+', ' ', text)  # normalize spaces/newlines
    return text.strip()

def chunk_text(text, chunk_size=300):
    words = text.split()
    print(f"Total words: {len(words)}")
    chunks = []
    for i in range(0, len(words), chunk_size):
        chunk = " ".join(words[i:i+chunk_size])
        chunks.append(chunk)
    print(f"✅ Total chunks created: {len(chunks)}")
    return chunks